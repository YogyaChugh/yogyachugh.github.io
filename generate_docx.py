import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document()

# Set margins (0.45 inch on all sides)
for section in doc.sections:
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    # Total width is 8.5, margins take 0.9, so content width is 7.6 inches

# Helper to add bottom border
def add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

# Default style helper
def set_style(paragraph, space_after=1, space_before=0):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.line_spacing = 1.0

# Add Heading
def add_heading(text):
    p = doc.add_paragraph()
    set_style(p, space_after=1, space_before=3)
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    add_bottom_border(p)
    return p

# HEADING (Contact Info)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_style(p, space_after=2)
name = p.add_run('YOGYA CHUGH\n')
name.bold = True
name.font.name = 'Times New Roman'
name.font.size = Pt(16)

contact = p.add_run('Delhi, India | yogya.developer@gmail.com | +91 9650029959\nlinkedin.com/in/yogyachugh | github.com/YogyaChugh')
contact.font.name = 'Times New Roman'
contact.font.size = Pt(10)

# TECHNICAL SKILLS
add_heading('Technical Skills')
skills = {
    'Languages': 'C++, Rust, Python, TypeScript, JavaScript, SQL',
    'Core CS': 'Data Structures & Algorithms, OOP, DBMS, Computer Networks, Compiler Design',
    'Backend': 'Node.js, Hono, FastAPI, Flask, Django, REST API Design, JWT, OAuth 2.0',
    'Databases & Cloud': 'PostgreSQL, Supabase, Firebase/GCP, Cloudflare R2',
    'Frontend': 'React, React Native, Expo, HTML/CSS',
    'Tools & Platforms': 'Docker, Git, Linux'
}
for k, v in skills.items():
    p = doc.add_paragraph()
    set_style(p, space_after=1)
    p.paragraph_format.left_indent = Pt(0)
    run_k = p.add_run(k + ': ')
    run_k.bold = True
    run_k.font.name = 'Times New Roman'
    run_k.font.size = Pt(10)
    run_v = p.add_run(v)
    run_v.font.name = 'Times New Roman'
    run_v.font.size = Pt(10)

# WORK EXPERIENCE
add_heading('Work Experience')
def add_exp(title, loc, role, date, items):
    p = doc.add_paragraph()
    set_style(p, space_after=0)
    # Add tab stop at 7.6 inches
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)
    
    r1 = p.add_run(f'{title}')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    
    rtab = p.add_run(f'\t{loc}')
    rtab.font.name = 'Times New Roman'
    rtab.font.size = Pt(10)
    
    p2 = doc.add_paragraph()
    set_style(p2, space_after=1)
    tab_stops2 = p2.paragraph_format.tab_stops
    tab_stops2.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)
    
    r2 = p2.add_run(f'{role}')
    r2.italic = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(10)
    
    r2tab = p2.add_run(f'\t{date}')
    r2tab.italic = True
    r2tab.font.name = 'Times New Roman'
    r2tab.font.size = Pt(10)
    
    for i in items:
        p_item = doc.add_paragraph(style='List Bullet')
        set_style(p_item, space_after=1)
        p_item.paragraph_format.left_indent = Inches(0.25)
        ri = p_item.add_run(i)
        ri.font.name = 'Times New Roman'
        ri.font.size = Pt(10)

add_exp('Vardhman Electricals', 'Remote', 'Software Developer - Intern', 'Nov 2025 - Jun 2026', [
    'TypeScript/Hono REST API on Render - Google OAuth, JWT auth, SSE real-time sync serving 200+ concurrent users across 3+ admin dashboards; live at vardhmansaathi.shop',
    'Designed and deployed a normalized PostgreSQL schema on Supabase to back authentication, orders, and admin-dashboard data',
    'Electron + React admin desktop app and a React Native (Expo) Android app deployed on Google Play Store',
    'Cloudflare R2 for image storage; total hosting at $7/month'
])

add_exp('Client Application Project', 'Remote', 'Software Developer - Freelance', 'Jun 2026 - Present', [
    'Cross-platform iOS/Android app (React Native, Expo) to be deployed at Play Store and App Store',
    'PostgreSQL database on Supabase with Google/Apple auth, RLS, and row-level access policies',
    'Subscription based Payment Handling using Superwall'
])

# PROJECTS
add_heading('Projects')
def add_proj(title, tech, items, link=None):
    p = doc.add_paragraph()
    set_style(p, space_after=0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)
    
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    if link:
        rl = p.add_run(f' | {link}')
        rl.font.name = 'Times New Roman'
        rl.font.size = Pt(10)
    
    rtab = p.add_run(f'\t{tech}')
    rtab.italic = True
    rtab.font.name = 'Times New Roman'
    rtab.font.size = Pt(10)
    
    for i in items:
        p_item = doc.add_paragraph(style='List Bullet')
        set_style(p_item, space_after=1)
        p_item.paragraph_format.left_indent = Inches(0.25)
        ri = p_item.add_run(i)
        ri.font.name = 'Times New Roman'
        ri.font.size = Pt(10)

add_proj('Webelo - Browser Rendering Engine', 'C++', [
    'C++ browser rendering engine; DOM tree follows the WHATWG spec with preorder depth-first traversal',
    'v2.0 ships an HTML Visualizer with Python bindings and hosted docs; CSSOM and Render Tree stages are not yet implemented'
], 'github.com/YogyaChugh/Webelo')

add_proj('AnyStudio - No-Code Android App Builder', 'In Development (Rust)', [
    'No-code Android app builder in Rust - drag-and-drop UI canvas that exports a signed, installable APK',
    'Custom compiler generates Android bytecode from scratch, bypassing the Java toolchain entirely',
    'Handles backend logic using drag-and-drop blocks, with automatic dependency management (SDKs, apksigner, zipalign) and virtual device testing'
], 'Private Repository')

add_proj('SIH - Blockchain Travel Security System (Safario)', '2024 - 2025', [
    'Decentralized tourist identity app - verifiable digital ID cards backed by a blockchain, scannable via QR',
    'Built on React, ethers.js, and WalletConnect; includes live geolocation and English/Hindi support'
])

p_item = doc.add_paragraph(style='List Bullet')
set_style(p_item, space_after=1)
p_item.paragraph_format.left_indent = Inches(0.25)
r1 = p_item.add_run('Open Source (Django): ')
r1.bold = True
r1.font.name = 'Times New Roman'
r1.font.size = Pt(10)
r2 = p_item.add_run('Fixed a bug where an unresolved relative {% extends %} template path raised an unrelated AttributeError instead of a clear TemplateSyntaxError; merged as django/django #19295')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(10)

# EDUCATION
add_heading('Education')
p = doc.add_paragraph()
set_style(p, space_after=0)
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)

r1 = p.add_run('Panipat Institute of Engineering and Technology (KUK)')
r1.bold = True
r1.font.name = 'Times New Roman'
r1.font.size = Pt(11)

rtab = p.add_run('\tPanipat, India')
rtab.font.name = 'Times New Roman'
rtab.font.size = Pt(10)

p2 = doc.add_paragraph()
set_style(p2, space_after=1)
tab_stops2 = p2.paragraph_format.tab_stops
tab_stops2.add_tab_stop(Inches(7.6), WD_TAB_ALIGNMENT.RIGHT)

r2 = p2.add_run('B.Tech - Computer Science and Engineering')
r2.italic = True
r2.font.name = 'Times New Roman'
r2.font.size = Pt(10)

r2tab = p2.add_run('\t2023 - 2027')
r2tab.italic = True
r2tab.font.name = 'Times New Roman'
r2tab.font.size = Pt(10)

p_item = doc.add_paragraph(style='List Bullet')
set_style(p_item, space_after=1)
p_item.paragraph_format.left_indent = Inches(0.25)
ri = p_item.add_run('CGPA: 8.05/10')
ri.font.name = 'Times New Roman'
ri.font.size = Pt(10)

# ACHIEVEMENTS
add_heading('Achievements')
achievements = [
    'Built 4 projects (450+ hrs) for Hack Club Summer of Making, GitHub-sponsored - awarded a $1,649 Gigabyte Aero laptop (RTX 5070, Ryzen AI 7, 32 GB, 1 TB SSD)',
    '3rd place, Cyber AI Hackathon 2025 (University of Derby, UK)',
    '3rd place, Agentic AI Hackathon 2025 (Ulster University, UK)',
    'Selected in SIH internal rounds in both 2024 and 2025',
    'Founder & Lead, Technical Club - organized Frontend Frenzy, website development competition',
    '1st place, Pattern Coding Competition 2023 - Rs. 2,000 cash prize'
]
for a in achievements:
    p = doc.add_paragraph(style='List Bullet')
    set_style(p, space_after=1)
    p.paragraph_format.left_indent = Inches(0.25)
    ri = p.add_run(a)
    ri.font.name = 'Times New Roman'
    ri.font.size = Pt(10)

# INTERESTS
add_heading('Interests')
p = doc.add_paragraph()
set_style(p, space_after=0)
ri = p.add_run('Competitive Programming | Game Development | Neural Networks')
ri.font.name = 'Times New Roman'
ri.font.size = Pt(10)

doc.save('Yogya_Chugh_Resume_V7.docx')
print("Successfully generated Yogya_Chugh_Resume_V7.docx")
